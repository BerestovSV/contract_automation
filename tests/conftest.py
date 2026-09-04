"""Общие фикстуры тестов.

Все данные — синтетические. Реальные персональные, юридические и банковские
сведения в тестах не используются.
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Callable, Iterable, List, Optional, Sequence

import openpyxl
import pytest
from docx import Document


@pytest.fixture(autouse=True)
def isolated_app_home(tmp_path_factory, monkeypatch):
    """Изолирует настройки/реестр/журнал от каталога реального пользователя."""
    home = tmp_path_factory.mktemp("app_home")
    monkeypatch.setenv("CONTRACT_GENERATOR_HOME", str(home))
    yield home


@pytest.fixture
def make_xlsx(tmp_path) -> Callable[..., Path]:
    """Создаёт синтетическую карточку компании.

    ``sheets`` — список ``(имя_листа, строки, состояние)``, где строки —
    последовательности значений ячеек.
    """

    def _make(
        rows: Optional[Sequence[Sequence[object]]] = None,
        name: str = "card.xlsx",
        sheets: Optional[Sequence[tuple]] = None,
        number_formats: Optional[dict] = None,
    ) -> Path:
        workbook = openpyxl.Workbook()
        default = workbook.active

        pages = list(sheets) if sheets is not None else [("Карточка", rows or [], "visible")]

        for index, page in enumerate(pages):
            title, page_rows, *rest = page
            state = rest[0] if rest else "visible"
            sheet = default if index == 0 else workbook.create_sheet()
            sheet.title = title
            sheet.sheet_state = state
            for row in page_rows:
                sheet.append(list(row))

        if number_formats:
            for (sheet_name, coordinate), fmt in number_formats.items():
                workbook[sheet_name][coordinate].number_format = fmt

        path = tmp_path / name
        workbook.save(path)
        workbook.close()
        return path

    return _make


@pytest.fixture
def base_rows() -> List[List[object]]:
    """Типовая карточка компании (синтетические данные)."""
    return [
        ["Форма собственности", "ООО"],
        ["Наименование компании", 'ООО «Ромашка»'],
        ["Место нахождения (юридический адрес)", "123456, г. Москва, ул. Тестовая, д. 1"],
        ["Почтовый адрес", "123456, г. Москва, а/я 1"],
        ["ОГРН", "1027700132195"],
        ["ИНН", "7707083893"],
        ["КПП", "770701001"],
        ["р/с", "40702810900000005555"],
        ["к/с", "30101810400000000225"],
        ["Банк", "ТЕСТБАНК"],
        ["БИК", "044525225"],
        ["Тел", "+7 900 000-00-00"],
        ["e-mail", "test@example.com"],
        ["Подписант, в лице", "Иванов Иван Иванович"],
        ["Подписант, кратко", "И.И. Иванов"],
        ["Должность подписанта", "Генеральный директор"],
        ["Пол подписанта", "мужской"],
        ["Действует на основании", "Устава"],
        ["Номер договора", "СЛД-0001-2025-М"],
        ["Дата договора", "01.03.2025"],
        ["Срок действия, в годах", 1],
        ["Дата окончания действия договора", "01.03.2026"],
    ]


@pytest.fixture
def make_docx(tmp_path) -> Callable[..., Path]:
    """Создаёт синтетический шаблон .docx.

    ``paragraphs`` — список либо строк, либо списков «кусков» текста; список
    кусков превращается в несколько run-ов (имитация разбиения Word).
    """

    def _make(
        paragraphs: Iterable = (),
        name: str = "template.docx",
        table_cells: Optional[Sequence[Sequence[str]]] = None,
        header: Optional[str] = None,
        footer: Optional[str] = None,
        nested_table_text: Optional[str] = None,
    ) -> Path:
        doc = Document()

        for item in paragraphs:
            paragraph = doc.add_paragraph()
            chunks = [item] if isinstance(item, str) else list(item)
            for chunk in chunks:
                if isinstance(chunk, tuple):
                    text, style = chunk
                else:
                    text, style = chunk, {}
                run = paragraph.add_run(text)
                for attribute, value in style.items():
                    setattr(run, attribute, value)

        if table_cells:
            table = doc.add_table(rows=len(table_cells), cols=len(table_cells[0]))
            for r, row in enumerate(table_cells):
                for c, text in enumerate(row):
                    table.cell(r, c).paragraphs[0].add_run(text)

            if nested_table_text is not None:
                inner = table.cell(0, 0).add_table(rows=1, cols=1)
                inner.cell(0, 0).paragraphs[0].add_run(nested_table_text)

        if header is not None:
            doc.sections[0].header.paragraphs[0].add_run(header)
        if footer is not None:
            doc.sections[0].footer.paragraphs[0].add_run(footer)

        path = tmp_path / name
        doc.save(path)
        return path

    return _make


def docx_text(path) -> str:
    """Весь текст документа, включая таблицы и колонтитулы."""
    from contract_generator.docx_filler import iter_all_paragraphs, paragraph_text

    doc = Document(str(path))
    return "\n".join(paragraph_text(p) for p in iter_all_paragraphs(doc))
