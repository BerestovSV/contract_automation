"""Сквозной интеграционный тест на синтетических файлах.

Проверяет весь путь: карточка .xlsx -> проверка -> шаблон .docx с разбитыми
плейсхолдерами -> готовый договор, включая точный формат даты окончания и
отсутствие SQLite-реестра.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

import openpyxl
import pytest
from docx import Document

from conftest import docx_text
from contract_generator.docx_filler import (
    PlaceholderError,
    fill_template,
    find_placeholders_in_file,
)
from contract_generator.excel_reader import read_company_card
from contract_generator.service import build_context, validate_card

MANUAL_NUMBER = "СЛД-2025/0042-М"
EXPECTED_END_DATE = '"04" сентября 2029 г.'


@pytest.fixture
def card_file(tmp_path) -> Path:
    """Синтетическая карточка компании (реальных данных нет)."""
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Карточка"
    rows = [
        ["Форма собственности", "ООО"],
        ["Наименование организации", 'ООО «Ромашка»'],
        ["Юр. адрес", "123456, г. Москва, ул. Тестовая, д. 1"],
        ["ИНН организации", 7707083893],          # число, не строка
        ["КПП организации", 770701001],
        ["ОГРН", "1027700132195"],
        ["БИК", "044525225"],
        ["Банк", "ТЕСТБАНК"],
        ["Расчетный счет", "40702810900000005555"],
        ["Корреспондентский счет", "30101810400000000225"],
        ["Электронная почта", "test@example.com"],
        ["Телефон", "+7 900 000-00-00"],
        ["ФИО подписанта", "Петрова Анна Сергеевна"],
        ["Пол подписанта", "женский"],            # явно указан
        ["Должность", "Генеральный директор"],
        ["Основание", "Устава"],
        ["Дата договора", "01.03.2025"],
        # Нативная ячейка-дата Excel.
        ["Дата окончания действия договора", datetime(2029, 9, 4)],
    ]
    for row in rows:
        sheet.append(row)
    path = tmp_path / "карточка.xlsx"
    workbook.save(path)
    workbook.close()
    return path


@pytest.fixture
def template_file(tmp_path) -> Path:
    """Шаблон с плейсхолдерами, разбитыми на несколько run-ов."""
    doc = Document()

    # {contract_number} и {contract_end_date} разбиты, как это делает Word.
    first = doc.add_paragraph()
    for chunk in ["Договор № {contract_num", "ber} от {contract_date},"]:
        first.add_run(chunk)

    second = doc.add_paragraph()
    second.add_run("действует до ")
    for chunk in ["{contract_end", "_date}", "."]:
        second.add_run(chunk)

    third = doc.add_paragraph()
    third.add_run("{company_name_full}, ИНН ").bold = True
    third.add_run("{inn}, КПП {kpp}, {legal_address}.")

    fourth = doc.add_paragraph()
    fourth.add_run("в лице {signatory_position_genitive} {signatory_full}, ")
    fourth.add_run("{acting_form_full}.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Банк"
    table.cell(0, 1).text = "{bank_name}, БИК {bik}"
    table.cell(1, 0).text = "Счёт"
    table.cell(1, 1).text = "{bank_account}"

    doc.sections[0].header.paragraphs[0].add_run("{company_name_short}")
    doc.sections[0].footer.paragraphs[0].add_run("Договор {contract_number}")

    path = tmp_path / "шаблон.docx"
    doc.save(path)
    return path


def test_end_to_end_contract_generation(card_file, template_file, tmp_path, isolated_app_home):
    template_before = template_file.read_bytes()
    card_before = card_file.read_bytes()

    # 1-5. Читаем карточку, вводим номер вручную, проверяем.
    card = read_company_card(card_file)
    card.set("contract_number", f"  {MANUAL_NUMBER}  ")  # с лишними пробелами

    result = validate_card(card, template_file)
    assert not result.is_blocked, [str(i) for i in result.errors]

    # 6. Генерация.
    context = build_context(card).context
    output = tmp_path / "готовый_договор.docx"
    report = fill_template(template_file, context, output)

    assert report.success
    assert output.exists()
    text = docx_text(output)

    # 7. Точный формат даты окончания.
    assert EXPECTED_END_DATE in text
    assert "«04»" not in text

    # 8. Ни одного выражения {...} не осталось.
    assert find_placeholders_in_file(output) == set()
    assert "{" not in text and "}" not in text

    # Номер обрезан только по краям.
    assert f"Договор № {MANUAL_NUMBER} от" in text
    assert "Договор № СЛД-2025/0042-М от «01» марта 2025 г." in text

    # Прочие подстановки.
    assert "Общество с ограниченной ответственностью «Ромашка»" in text
    assert "ИНН 7707083893, КПП 770701001" in text       # число не стало 7.7e+09
    assert "Генерального директора Петрова Анна Сергеевна" in text
    assert "действующей на основании Устава" in text     # женский род
    assert "ТЕСТБАНК, БИК 044525225" in text             # таблица
    assert "Ромашка" in text                             # колонтитул

    # 9. Исходные файлы не изменены.
    assert template_file.read_bytes() == template_before
    assert card_file.read_bytes() == card_before

    # 10. Никакого SQLite-реестра.
    assert not list(Path(isolated_app_home).rglob("*.sqlite3"))
    assert not list(tmp_path.rglob("*.sqlite3"))


def test_end_to_end_blocked_by_empty_field(card_file, template_file, tmp_path):
    """Пустое поле-источник останавливает весь конвейер."""
    card = read_company_card(card_file)
    card.set("contract_number", MANUAL_NUMBER)
    card.set("bank_name", "")

    result = validate_card(card, template_file)
    assert result.is_blocked
    assert any(i.field == "bank_name" for i in result.errors)

    output = tmp_path / "не_должен_появиться.docx"
    with pytest.raises(PlaceholderError) as exc:
        fill_template(template_file, build_context(card).context, output)

    assert "{bank_name}" in exc.value.report.empty_values
    assert not output.exists()
    assert not list(tmp_path.glob(".contract_*"))


def test_end_to_end_blocked_by_unknown_placeholder(card_file, tmp_path):
    doc = Document()
    doc.add_paragraph("{inn} и {contract_namber}")
    template = tmp_path / "с_опечаткой.docx"
    doc.save(template)

    card = read_company_card(card_file)
    card.set("contract_number", MANUAL_NUMBER)

    result = validate_card(card, template)
    assert result.is_blocked
    assert any("{contract_namber}" in i.message for i in result.errors)

    output = tmp_path / "нет.docx"
    with pytest.raises(PlaceholderError):
        fill_template(template, build_context(card).context, output)
    assert not output.exists()
