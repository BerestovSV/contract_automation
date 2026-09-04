"""Тесты замены плейсхолдеров в DOCX."""
from __future__ import annotations

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

from conftest import docx_text
from contract_generator.docx_filler import (
    TemplateError,
    fill_template,
    find_placeholders_in_file,
    paragraph_text,
    replace_in_paragraph,
    safe_filename,
    unique_path,
)

CTX = {"inn": "7707083893", "company_name": "ООО «Ромашка»", "kpp": "770701001"}


def test_placeholder_in_single_run(make_docx, tmp_path):
    template = make_docx(["ИНН: {inn}"])
    out = tmp_path / "out.docx"
    fill_template(template, CTX, out)
    assert "ИНН: 7707083893" in docx_text(out)


def test_placeholder_split_across_runs(make_docx, tmp_path):
    template = make_docx([["ИНН: {i", "n", "n", "} конец"]])
    out = tmp_path / "out.docx"
    report = fill_template(template, CTX, out)
    text = docx_text(out)
    assert "ИНН: 7707083893 конец" in text
    assert report.replaced == {"{inn}": 1}
    assert not report.remaining_placeholders


def test_placeholder_split_with_braces_alone(make_docx, tmp_path):
    template = make_docx([["{", "inn", "}"]])
    out = tmp_path / "out.docx"
    fill_template(template, CTX, out)
    assert "7707083893" in docx_text(out)


def test_multiple_placeholders_in_one_paragraph(make_docx, tmp_path):
    template = make_docx([["{company_name}", ", ИНН {i", "nn}, КПП {kpp}"]])
    out = tmp_path / "out.docx"
    report = fill_template(template, CTX, out)
    assert "ООО «Ромашка», ИНН 7707083893, КПП 770701001" in docx_text(out)
    assert set(report.replaced) == {"{company_name}", "{inn}", "{kpp}"}


def test_placeholders_in_tables_and_nested_tables(make_docx, tmp_path):
    template = make_docx(
        ["шапка"],
        table_cells=[["ИНН", "{inn}"], ["КПП", "{kpp}"]],
        nested_table_text="{company_name}",
    )
    out = tmp_path / "out.docx"
    fill_template(template, CTX, out)
    text = docx_text(out)
    assert "7707083893" in text
    assert "770701001" in text
    assert "ООО «Ромашка»" in text


def test_placeholders_in_header_and_footer(make_docx, tmp_path):
    template = make_docx(["тело"], header="Шапка {inn}", footer="Подвал {kpp}")
    out = tmp_path / "out.docx"
    fill_template(template, CTX, out)
    text = docx_text(out)
    assert "Шапка 7707083893" in text
    assert "Подвал 770701001" in text


def test_formatting_preserved_around_replacement(make_docx, tmp_path):
    template = make_docx([[
        ("Жирное ", {"bold": True}),
        ("{inn}", {"italic": True}),
        (" курсив", {"italic": True}),
    ]])
    out = tmp_path / "out.docx"
    fill_template(template, CTX, out)

    doc = Document(str(out))
    runs = doc.paragraphs[0].runs
    assert runs[0].text == "Жирное " and runs[0].bold is True
    replaced = next(r for r in runs if "7707083893" in r.text)
    assert replaced.italic is True
    assert any(r.text == " курсив" and r.italic for r in runs)


def test_formatting_of_untouched_runs_preserved_when_split(make_docx, tmp_path):
    template = make_docx([[
        ("Начало ", {"bold": True}),
        ("{in", {"italic": True}),
        ("n}", {"italic": True}),
        (" хвост", {"underline": True}),
    ]])
    out = tmp_path / "out.docx"
    fill_template(template, CTX, out)

    doc = Document(str(out))
    runs = doc.paragraphs[0].runs
    assert runs[0].bold is True and runs[0].text == "Начало "
    tail = next(r for r in runs if r.text == " хвост")
    assert tail.underline is True
    assert "7707083893" in paragraph_text(doc.paragraphs[0])


def test_font_size_and_color_preserved(make_docx, tmp_path):
    from docx import Document as D

    doc = D()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("{inn}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    template = tmp_path / "t.docx"
    doc.save(template)

    out = tmp_path / "out.docx"
    fill_template(template, CTX, out)

    result = D(str(out)).paragraphs[0].runs[0]
    assert result.text == "7707083893"
    assert result.font.size == Pt(14)
    assert result.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)


def test_unknown_placeholder_reported_but_not_blocking(make_docx, tmp_path):
    template = make_docx(["{inn} и {nonexistent_field}"])
    out = tmp_path / "out.docx"
    report = fill_template(template, CTX, out)
    assert report.unknown_placeholders == ["{nonexistent_field}"]
    assert report.success
    assert out.exists()


def test_empty_value_reported(make_docx, tmp_path):
    template = make_docx(["Телефон: {phone}"])
    out = tmp_path / "out.docx"
    report = fill_template(template, {**CTX, "phone": ""}, out)
    assert "{phone}" in report.empty_values
    assert "Телефон: " in docx_text(out)


def test_required_placeholder_left_unresolved_blocks_output(make_docx, tmp_path):
    """Если требуемый плейсхолдер остался, файл не создаётся."""
    template = make_docx(["{inn} и {contract_number}"])
    out = tmp_path / "out.docx"

    with pytest.raises(TemplateError):
        fill_template(template, CTX, out, required_placeholders=["{contract_number}"])

    assert not out.exists()
    # Временные файлы удалены.
    assert not list(tmp_path.glob(".contract_*"))


def test_source_template_not_modified(make_docx, tmp_path):
    template = make_docx(["{inn}"])
    before = template.read_bytes()
    fill_template(template, CTX, tmp_path / "out.docx")
    assert template.read_bytes() == before


def test_find_placeholders_discovers_split_ones(make_docx):
    template = make_docx([["{comp", "any_name} и {i", "nn}"]], header="{kpp}")
    assert find_placeholders_in_file(template) == {
        "{company_name}", "{inn}", "{kpp}"
    }


def test_replace_in_paragraph_returns_counts(make_docx):
    doc = Document(str(make_docx(["{inn} {inn}"])))
    counts = replace_in_paragraph(doc.paragraphs[0], {"{inn}": "X"})
    assert counts == {"{inn}": 2}
    assert paragraph_text(doc.paragraphs[0]) == "X X"


@pytest.mark.parametrize(
    "raw,expected",
    [
        ('ООО "Тест"/Договор', "ООО _Тест__Договор"),
        ("CON", "_CON"),
        ("   ", "contract"),
        ("имя<>:|?*", "имя______"),
    ],
)
def test_safe_filename(raw, expected):
    assert safe_filename(raw) == expected


def test_unique_path_avoids_collision(tmp_path):
    (tmp_path / "a.docx").write_bytes(b"x")
    assert unique_path(tmp_path, "a.docx").name == "a_1.docx"
    (tmp_path / "a_1.docx").write_bytes(b"x")
    assert unique_path(tmp_path, "a.docx").name == "a_2.docx"


def test_missing_template_raises(tmp_path):
    with pytest.raises(TemplateError):
        fill_template(tmp_path / "нет.docx", CTX, tmp_path / "o.docx")
