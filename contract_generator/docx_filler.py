"""Заполнение шаблона договора с сохранением форматирования.

Ключевая задача — заменить плейсхолдеры вида ``{field}``, которые Word почти
всегда разбивает на несколько ``run``-ов (из-за проверки орфографии, правок,
языковых меток). Наивная реализация склеивает параграф в один run и теряет
всё форматирование. Здесь применяется другой подход:

* текст параграфа собирается из всех его run-ов (включая run-ы внутри
  гиперссылок) с запоминанием границ;
* найденные вхождения заменяются «на месте»: значение целиком помещается в
  первый run вхождения (наследуя его формат), из остальных run-ов вхождения
  вырезаются только попавшие в него символы;
* run-ы, не затронутые заменой, не изменяются вообще.

Таким образом сохраняются жирность, курсив, подчёркивание, шрифт, размер,
цвет, гиперссылки и форматирование параграфа.

Строгое правило безопасности (договор — юридический документ): любой
неизвестный, пустой или оставшийся после замены плейсхолдер ОТМЕНЯЕТ
генерацию. Это поведение по умолчанию и его нельзя отключить параметром
вызова.
"""
from __future__ import annotations

import logging
import os
import re
import tempfile
from pathlib import Path
from typing import Dict, Iterator, List, Optional, Sequence, Set, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .models import GenerationReport
from .placeholders import PLACEHOLDER_RE, is_supported

logger = logging.getLogger(__name__)


class TemplateError(Exception):
    """Шаблон не удалось открыть или обработать."""


class PlaceholderError(TemplateError):
    """Генерация отменена: шаблон содержит незаполнимые плейсхолдеры.

    Исключение несёт структурированный :class:`GenerationReport`, чтобы
    интерфейс показывал подробности, не разбирая текст сообщения.
    """

    def __init__(self, report: GenerationReport) -> None:
        self.report = report
        super().__init__(report.failure_message_ru())


# --------------------------------------------------------------------------
# Обход документа
# --------------------------------------------------------------------------

def _iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell: _Cell) -> Iterator[Paragraph]:
    yield from cell.paragraphs
    for nested in cell.tables:  # вложенные таблицы
        yield from _iter_table_paragraphs(nested)


def _iter_container_paragraphs(container) -> Iterator[Paragraph]:
    """Параграфы тела документа, колонтитула или ячейки, включая таблицы."""
    yield from container.paragraphs
    for table in container.tables:
        yield from _iter_table_paragraphs(table)


_HEADER_FOOTER_ATTRS = (
    "header",
    "footer",
    "even_page_header",
    "even_page_footer",
    "first_page_header",
    "first_page_footer",
)


def iter_all_paragraphs(doc) -> Iterator[Paragraph]:
    """Все параграфы документа: тело, таблицы, вложенные таблицы, колонтитулы."""
    yield from _iter_container_paragraphs(doc)
    for section in doc.sections:
        for attr in _HEADER_FOOTER_ATTRS:
            part = getattr(section, attr, None)
            if part is None:
                continue
            yield from _iter_container_paragraphs(part)


def _paragraph_runs(paragraph: Paragraph) -> List[Run]:
    """Все run-ы параграфа в порядке документа, включая run-ы гиперссылок.

    ``Paragraph.runs`` в python-docx не возвращает run-ы внутри ``w:hyperlink``,
    поэтому обходим XML напрямую.
    """
    return [
        Run(element, paragraph)
        for element in paragraph._p.iter(qn("w:r"))
    ]


# --------------------------------------------------------------------------
# Поиск плейсхолдеров
# --------------------------------------------------------------------------

def paragraph_text(paragraph: Paragraph) -> str:
    """Полный текст параграфа, собранный из run-ов (в т.ч. в гиперссылках)."""
    return "".join(run.text for run in _paragraph_runs(paragraph))


def find_placeholders(doc) -> Set[str]:
    """Все выражения ``{...}`` в документе."""
    found: Set[str] = set()
    for paragraph in iter_all_paragraphs(doc):
        found.update(PLACEHOLDER_RE.findall(paragraph_text(paragraph)))
    return found


def find_placeholders_in_file(path: str | Path) -> Set[str]:
    """Все выражения ``{...}`` в файле ``.docx``."""
    try:
        doc = Document(str(path))
    except Exception as exc:
        raise TemplateError(f"Не удалось открыть документ: {exc}") from exc
    return find_placeholders(doc)


# --------------------------------------------------------------------------
# Замена
# --------------------------------------------------------------------------

def replace_in_paragraph(paragraph: Paragraph, replacements: Dict[str, str]) -> Dict[str, int]:
    """Заменяет плейсхолдеры в параграфе, сохраняя форматирование run-ов.

    Возвращает ``{плейсхолдер: количество замен}``.
    """
    runs = _paragraph_runs(paragraph)
    if not runs:
        return {}

    texts = [run.text for run in runs]
    full_text = "".join(texts)
    if "{" not in full_text:
        return {}

    # Границы run-ов: offsets[i] — позиция начала run i в full_text.
    offsets: List[int] = []
    position = 0
    for text in texts:
        offsets.append(position)
        position += len(text)

    matches: List[Tuple[int, int, str, str]] = []
    for match in PLACEHOLDER_RE.finditer(full_text):
        token = match.group(0)
        if token in replacements:
            matches.append((match.start(), match.end(), token, replacements[token]))

    if not matches:
        return {}

    counts: Dict[str, int] = {}
    new_texts = list(texts)

    # Идём с конца, чтобы смещения ранее найденных вхождений оставались верными.
    for start, end, token, value in reversed(matches):
        first = _run_index_at(offsets, texts, start)
        last = _run_index_at(offsets, texts, end - 1)

        if first == last:
            local_start = start - offsets[first]
            local_end = end - offsets[first]
            new_texts[first] = (
                new_texts[first][:local_start] + value + new_texts[first][local_end:]
            )
        else:
            # Значение целиком уходит в первый run вхождения — он задаёт формат.
            new_texts[first] = new_texts[first][: start - offsets[first]] + value
            for middle in range(first + 1, last):
                new_texts[middle] = ""
            new_texts[last] = new_texts[last][end - offsets[last]:]

        counts[token] = counts.get(token, 0) + 1

    for run, old, new in zip(runs, texts, new_texts):
        if old != new:
            run.text = new

    return counts


def _run_index_at(offsets: Sequence[int], texts: Sequence[str], position: int) -> int:
    """Индекс run-а, которому принадлежит символ ``position``."""
    for index in range(len(offsets) - 1, -1, -1):
        if texts[index] and offsets[index] <= position < offsets[index] + len(texts[index]):
            return index
    return 0


def replace_in_document(doc, replacements: Dict[str, str]) -> Dict[str, int]:
    """Заменяет плейсхолдеры во всём документе. Возвращает счётчик замен."""
    total: Dict[str, int] = {}
    for paragraph in iter_all_paragraphs(doc):
        for token, count in replace_in_paragraph(paragraph, replacements).items():
            total[token] = total.get(token, 0) + count
    return total


# --------------------------------------------------------------------------
# Генерация документа
# --------------------------------------------------------------------------

def build_replacements(context: Dict[str, str]) -> Dict[str, str]:
    """Строит таблицу замен ``{ключ}`` -> значение из контекста."""
    return {
        "{%s}" % key: "" if value is None else str(value)
        for key, value in context.items()
    }


def fill_template(
    template_path: str | Path,
    context: Dict[str, str],
    output_path: str | Path,
) -> GenerationReport:
    """Создаёт заполненный договор из шаблона.

    Безопасность не зависит от параметров вызова: функция сама находит все
    плейсхолдеры шаблона и отменяет генерацию, если хотя бы один из них
    неизвестен, пуст или остался в сохранённом документе.

    Порядок работы:

    1. найти все плейсхолдеры шаблона;
    2. убедиться, что каждый из них известен приложению;
    3. получить значение каждого из ``context``;
    4. убедиться, что значение непустое;
    5. при любой неудаче — возбудить :class:`PlaceholderError`, не создавая
       выходной файл;
    6. сохранить результат во временный файл, повторно открыть его и убедиться,
       что выражений ``{...}`` не осталось;
    7. только после этого перенести временный файл на место назначения.

    :raises PlaceholderError: шаблон содержит незаполнимые плейсхолдеры;
        отчёт доступен в атрибуте ``report``.
    :raises TemplateError: шаблон не найден или не читается.
    """
    template = Path(template_path)
    output = Path(output_path)
    report = GenerationReport(template_path=str(template), output_path=str(output))

    if not template.is_file():
        raise TemplateError(f"Файл шаблона не найден: {template}")

    try:
        doc = Document(str(template))
    except Exception as exc:
        logger.exception("Не удалось открыть шаблон")
        raise TemplateError(
            f"Не удалось открыть файл шаблона «{template.name}»: {exc}"
        ) from exc

    replacements = build_replacements(context)

    found = find_placeholders(doc)
    report.template_placeholders = sorted(found)

    # 2. Неизвестные плейсхолдеры — блокирующая ошибка.
    report.unknown_placeholders = sorted(
        token for token in found
        if not is_supported(token) or token not in replacements
    )

    # 3-4. Известные плейсхолдеры с пустым значением — блокирующая ошибка.
    report.empty_values = sorted(
        token for token in found
        if token not in report.unknown_placeholders
        and not replacements.get(token, "").strip()
    )

    if report.unknown_placeholders or report.empty_values:
        logger.error(
            "Генерация отменена: неизвестных плейсхолдеров %d, пустых %d",
            len(report.unknown_placeholders), len(report.empty_values),
        )
        raise PlaceholderError(report)

    report.replaced = replace_in_document(doc, replacements)

    output.parent.mkdir(parents=True, exist_ok=True)
    tmp_handle, tmp_name = tempfile.mkstemp(
        prefix=".contract_", suffix=".docx", dir=str(output.parent)
    )
    os.close(tmp_handle)
    tmp_path = Path(tmp_name)

    try:
        doc.save(str(tmp_path))

        # 6. Контроль по фактически сохранённому файлу: остаться не должно
        #    НИ ОДНОГО выражения {...}, включая незамеченные ранее.
        report.remaining_placeholders = sorted(find_placeholders_in_file(tmp_path))
        if report.remaining_placeholders:
            logger.error(
                "Генерация отменена: в сохранённом файле остались плейсхолдеры %s",
                report.remaining_placeholders,
            )
            raise PlaceholderError(report)

        # 7. Перенос на место назначения — только после успешной проверки.
        os.replace(str(tmp_path), str(output))
    finally:
        _remove_quietly(tmp_path)

    logger.info(
        "Договор сформирован: %s (замен: %d)",
        output.name,
        sum(report.replaced.values()),
    )
    return report


def _remove_quietly(path: Path) -> None:
    try:
        if path.exists():
            path.unlink()
    except OSError:
        logger.warning("Не удалось удалить временный файл %s", path)


_UNSAFE_CHARS_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
_RESERVED_NAMES = {
    "CON", "PRN", "AUX", "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}


def safe_filename(name: str, default: str = "contract", max_length: int = 120) -> str:
    """Делает строку безопасной для имени файла в Windows."""
    cleaned = _UNSAFE_CHARS_RE.sub("_", str(name or "")).strip()
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    if not cleaned:
        return default
    if cleaned.split(".")[0].upper() in _RESERVED_NAMES:
        cleaned = "_" + cleaned
    return cleaned[:max_length].strip(" .") or default


def unique_path(directory: str | Path, filename: str) -> Path:
    """Возвращает несуществующий путь, добавляя ``_1``, ``_2``, ... при коллизии."""
    base = Path(directory)
    candidate = base / filename
    if not candidate.exists():
        return candidate
    stem, suffix = candidate.stem, candidate.suffix
    counter = 1
    while True:
        candidate = base / f"{stem}_{counter}{suffix}"
        if not candidate.exists():
            return candidate
        counter += 1
