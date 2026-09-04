"""
Логика заполнения шаблона договора - с поддержкой колонтитулов и обработкой ошибок
"""
import os
import re
from datetime import datetime
import openpyxl

# Добавляем текущую директорию в путь
import sys
current_dir = os.path.dirname(os.path.abspath(__file__))
if current_dir not in sys.path:
    sys.path.insert(0, current_dir)

try:
    from config import (
        MONTHS_RU,
        POSITION_CASES,
        OWNERSHIP_CASES,
        OWNERSHIP_FULL_NAMES,
        GENDER_AGREEMENT
    )
except ImportError:
    # Если config не найден, используем значения по умолчанию
    MONTHS_RU = {
        '01': 'января', '02': 'февраля', '03': 'марта',
        '04': 'апреля', '05': 'мая', '06': 'июня',
        '07': 'июля', '08': 'августа', '09': 'сентября',
        '10': 'октября', '11': 'ноября', '12': 'декабря'
    }
    POSITION_CASES = {}
    OWNERSHIP_CASES = {}
    OWNERSHIP_FULL_NAMES = {}
    GENDER_AGREEMENT = {
        'masculine': {
            'based_on': 'действующего на основании',
            'based_on_full': 'действующего на основании {based_on}',
            'acting': 'действующего',
            'acting_full': 'действующего на основании {based_on}',
        },
        'feminine': {
            'based_on': 'действующей на основании',
            'based_on_full': 'действующей на основании {based_on}',
            'acting': 'действующей',
            'acting_full': 'действующей на основании {based_on}',
        }
    }

# Пробуем импортировать docx с обработкой ошибок
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.section import WD_SECTION
except ImportError as e:
    print(f"Ошибка импорта python-docx: {e}")
    print("Установите: pip install python-docx==1.1.0")
    raise



def load_company_data(excel_file):
    """
    Загрузка данных компании из Excel файла
    """
    try:
        wb = openpyxl.load_workbook(excel_file, data_only=True)
        ws = wb.active
        
        data = {}
        for row in ws.iter_rows(min_row=1, max_col=2, values_only=True):
            if row[0] and row[1]:
                field = str(row[0]).strip().replace(':', '').strip()
                value = str(row[1]).strip() if row[1] else ''
                data[field] = value
        
        wb.close()
        return data
    except Exception as e:
        print(f"Ошибка загрузки Excel: {e}")
        raise


def detect_gender(signatory_name):
    """
    Определение пола подписанта по имени и отчеству
    """
    if not signatory_name:
        return 'masculine'
    
    feminine_patterns = [
        r'на$',
        r'ая$',
        r'ья$',
        r'ия$',
    ]
    
    name_parts = signatory_name.strip().split()
    if name_parts:
        first_name = name_parts[0]
        for pattern in feminine_patterns:
            if re.search(pattern, first_name, re.IGNORECASE):
                return 'feminine'
    
    if len(name_parts) >= 3:
        patronymic = name_parts[2]
        if patronymic.endswith('на') or patronymic.endswith('НА'):
            return 'feminine'
    
    return 'masculine'


def decline_ownership(ownership_form, case):
    """
    Склонение формы собственности по падежам
    """
    if not ownership_form:
        return ''
    
    ownership_form = ownership_form.strip().strip('"').strip("'")
    
    if case in OWNERSHIP_CASES:
        case_dict = OWNERSHIP_CASES[case]
        if ownership_form in case_dict:
            return case_dict[ownership_form]
    
    return ownership_form


def get_ownership_full_name(ownership_form):
    """
    Получение полного названия формы собственности
    """
    if not ownership_form:
        return ''
    
    ownership_form = ownership_form.strip().strip('"').strip("'")
    return OWNERSHIP_FULL_NAMES.get(ownership_form, ownership_form)


def decline_position(position, case):
    """
    Склонение должности по падежам
    """
    if not position:
        return position
    
    if case in POSITION_CASES:
        case_dict = POSITION_CASES[case]
        if position in case_dict:
            return case_dict[position]
    
    for base_position, declined in POSITION_CASES.get(case, {}).items():
        if base_position in position:
            return position.replace(base_position, declined)
    
    return position


def get_gender_agreement(gender, based_on, context='based_on'):
    """
    Получение согласованной формы в зависимости от пола
    """
    if gender not in GENDER_AGREEMENT:
        gender = 'masculine'
    
    if context == 'based_on':
        return GENDER_AGREEMENT[gender]['based_on']
    elif context == 'based_on_full':
        return GENDER_AGREEMENT[gender]['based_on_full'].replace('{based_on}', based_on)
    elif context == 'acting':
        return GENDER_AGREEMENT[gender]['acting']
    elif context == 'acting_full':
        return GENDER_AGREEMENT[gender]['acting_full'].replace('{based_on}', based_on)
    
    return GENDER_AGREEMENT[gender]['based_on']


def get_company_info(data):
    """
    Извлечение информации о компании из данных с преобразованием форматов
    """
    company_info = {}
    
    field_map = {
        'Форма собственности': 'ownership_form',
        'Наименование компании': 'company_name',
        'Место нахождения (юридический адрес)': 'legal_address',
        'Почтовый адрес': 'postal_address',
        'ОГРН': 'ogrn',
        'ИНН': 'inn',
        'КПП': 'kpp',
        'р/с': 'bank_account',
        'к/с': 'corr_account',
        'Банк': 'bank_name',
        'БИК': 'bik',
        'Тел': 'phone',
        'e-mail': 'email',
        'Подписант, в лице': 'signatory_full',
        'Подписант, кратко': 'signatory_short',
        'Должность подписанта': 'signatory_position',
        'Пол подписанта': 'signatory_gender',
        'Действует на основании': 'based_on',
        'ЭДО': 'edo_provider',
        'Домены заказчика': 'domains',
        'Номер договора': 'contract_number',
        'Дата договора': 'contract_date',
        'Срок действия, в годах': 'contract_term_years',
        'Дата окончания действия договора': 'contract_end_date',
    }
    
    for excel_field, doc_field in field_map.items():
        if excel_field in data:
            company_info[doc_field] = data[excel_field]
        else:
            company_info[doc_field] = ''
    
    # Определяем пол подписанта
    signatory_name = company_info.get('signatory_full', '')
    if company_info.get('signatory_gender'):
        gender_value = company_info.get('signatory_gender', '').lower()
        if 'жен' in gender_value or 'female' in gender_value or gender_value in ['ж', 'f']:
            gender = 'feminine'
        else:
            gender = 'masculine'
    else:
        gender = detect_gender(signatory_name)
    
    company_info['signatory_gender'] = gender
    
    # Получаем согласованные формы
    based_on = company_info.get('based_on', '')
    company_info['based_on_agreed'] = get_gender_agreement(gender, based_on, 'based_on')
    company_info['based_on_agreed_full'] = get_gender_agreement(gender, based_on, 'based_on_full')
    
    if gender == 'masculine':
        company_info['acting_form'] = 'действующего'
        company_info['acting_form_full'] = f'действующего на основании {based_on}'
    else:
        company_info['acting_form'] = 'действующей'
        company_info['acting_form_full'] = f'действующей на основании {based_on}'
    
    # Обработка формы собственности
    ownership = company_info.get('ownership_form', '')
    if ownership:
        company_info['ownership_full'] = get_ownership_full_name(ownership)
        company_info['ownership_form_genitive'] = decline_ownership(ownership, 'genitive')
        company_info['ownership_form_dative'] = decline_ownership(ownership, 'dative')
        company_info['ownership_form_accusative'] = decline_ownership(ownership, 'accusative')
        company_info['ownership_form_instrumental'] = decline_ownership(ownership, 'instrumental')
        company_info['ownership_form_prepositional'] = decline_ownership(ownership, 'prepositional')
        
        company_name = company_info.get('company_name', '')
        if company_name:
            # Очищаем название от формы собственности
            clean_name = company_name
            for form in ['ООО', 'АО', 'ПАО', 'ЗАО', 'ОАО', 'ИП', 'ГУП', 'МУП', 'ФГУП', 'НКО', 'ТСЖ', 'СНТ', 'ДНТ', 'КФХ', 'ПК', 'СПК']:
                clean_name = clean_name.replace(f'{form} «', '').replace(f'{form} "', '').replace(f'{form} ', '')
            
            clean_name = clean_name.strip().strip('«').strip('»').strip('"').strip("'")
            
            if company_info['ownership_full']:
                company_info['company_name_full'] = f'{company_info["ownership_full"]} «{clean_name}»'
            else:
                company_info['company_name_full'] = f'{ownership} «{clean_name}»'
            
            company_info['company_name_short'] = clean_name
        else:
            company_info['company_name_full'] = ownership
            company_info['company_name_short'] = ''
    else:
        company_name = company_info.get('company_name', '')
        company_info['company_name_full'] = company_name
        company_info['company_name_short'] = company_name
        company_info['ownership_full'] = ''
        company_info['ownership_form_genitive'] = ''
        company_info['ownership_form_dative'] = ''
        company_info['ownership_form_accusative'] = ''
        company_info['ownership_form_instrumental'] = ''
        company_info['ownership_form_prepositional'] = ''
    
    # Преобразуем даты
    if company_info.get('contract_date'):
        company_info['contract_date_formatted'] = format_date_ru(company_info['contract_date'])
    
    if company_info.get('contract_end_date'):
        company_info['contract_end_date_formatted'] = format_date_ru(company_info['contract_end_date'])
    
    # Склоняем должность
    position = company_info.get('signatory_position', '')
    if position:
        company_info['signatory_position_genitive'] = decline_position(position, 'genitive')
        company_info['signatory_position_dative'] = decline_position(position, 'dative')
        company_info['signatory_position_accusative'] = decline_position(position, 'accusative')
        company_info['signatory_position_instrumental'] = decline_position(position, 'instrumental')
        company_info['signatory_position_prepositional'] = decline_position(position, 'prepositional')
    else:
        company_info['signatory_position_genitive'] = ''
        company_info['signatory_position_dative'] = ''
        company_info['signatory_position_accusative'] = ''
        company_info['signatory_position_instrumental'] = ''
        company_info['signatory_position_prepositional'] = ''
    
    return company_info


def format_date_ru(date_str):
    """
    Преобразование даты из формата дд.мм.гггг в "дд месяц гггг"
    """
    if not date_str:
        return ''
    
    date_str = date_str.strip().strip('"').strip("'")
    
    patterns = [
        r'(\d{2})\.(\d{2})\.(\d{4})',
        r'(\d{2})/(\d{2})/(\d{4})',
        r'(\d{2})-(\d{2})-(\d{4})',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, date_str)
        if match:
            day, month, year = match.groups()
            month_name = MONTHS_RU.get(month, month)
            return f'«{day}» {month_name} {year} г.'
    
    return date_str


def safe_get_run_text(paragraph):
    """
    Безопасное получение текста из параграфа
    """
    try:
        return ''.join(run.text for run in paragraph.runs)
    except:
        return paragraph.text


def safe_set_run_text(paragraph, text):
    """
    Безопасная установка текста в параграф
    """
    try:
        if paragraph.runs:
            # Сохраняем форматирование
            first_run = paragraph.runs[0]
            font_name = getattr(first_run.font, 'name', None)
            font_size = getattr(first_run.font, 'size', None)
            bold = getattr(first_run, 'bold', None)
            italic = getattr(first_run, 'italic', None)
            
            # Очищаем все run'ы
            for run in paragraph.runs:
                run.text = ''
            
            # Создаем один run с текстом
            new_run = paragraph.runs[0]
            new_run.text = text
            
            # Восстанавливаем форматирование
            if font_name:
                try:
                    new_run.font.name = font_name
                except:
                    pass
            if font_size:
                try:
                    new_run.font.size = font_size
                except:
                    pass
            if bold is not None:
                try:
                    new_run.bold = bold
                except:
                    pass
            if italic is not None:
                try:
                    new_run.italic = italic
                except:
                    pass
        else:
            paragraph.text = text
    except Exception as e:
        # Если не удалось сохранить форматирование, просто заменяем текст
        try:
            paragraph.text = text
        except:
            pass


def replace_text_in_paragraph(paragraph, replacements):
    """
    Замена текста в одном параграфе с объединением run'ов
    """
    try:
        # Получаем полный текст параграфа
        full_text = safe_get_run_text(paragraph)
        
        # Заменяем текст
        text = full_text
        for placeholder, value in replacements.items():
            if placeholder in text:
                text = text.replace(placeholder, value)
        
        # Если текст изменился, обновляем параграф
        if text != full_text:
            safe_set_run_text(paragraph, text)
    except Exception as e:
        # Если произошла ошибка, пробуем простую замену
        try:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        except:
            pass


def replace_text_in_doc(doc, replacements):
    """
    Замена текста во всем документе (включая колонтитулы и подвалы)
    """
    # 1. Основное тело документа
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    # 2. Таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)
    
    # 3. Колонтитулы
    try:
        for section in doc.sections:
            # Верхний колонтитул
            try:
                header = section.header
                for paragraph in header.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, replacements)
            except:
                pass
            
            # Нижний колонтитул
            try:
                footer = section.footer
                for paragraph in footer.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, replacements)
            except:
                pass
    except:
        pass


def fill_template(template_path, company_data, output_path):
    """
    Заполнение шаблона данными компании (с поддержкой колонтитулов)
    """
    try:
        # Загружаем шаблон
        doc = Document(template_path)
    except Exception as e:
        raise Exception(f"Не удалось открыть файл шаблона: {str(e)}")
    
    # Подготавливаем данные для замены
    replacements = {
        # Форма собственности
        '{ownership_form}': company_data.get('ownership_form', ''),
        '{ownership_full}': company_data.get('ownership_full', ''),
        '{ownership_form_genitive}': company_data.get('ownership_form_genitive', ''),
        '{ownership_form_dative}': company_data.get('ownership_form_dative', ''),
        '{ownership_form_accusative}': company_data.get('ownership_form_accusative', ''),
        '{ownership_form_instrumental}': company_data.get('ownership_form_instrumental', ''),
        '{ownership_form_prepositional}': company_data.get('ownership_form_prepositional', ''),
        
        # Названия компании
        '{company_name}': company_data.get('company_name', ''),
        '{company_name_full}': company_data.get('company_name_full', ''),
        '{company_name_short}': company_data.get('company_name_short', ''),
        
        # Основные реквизиты
        '{ogrn}': company_data.get('ogrn', ''),
        '{inn}': company_data.get('inn', ''),
        '{kpp}': company_data.get('kpp', ''),
        '{legal_address}': company_data.get('legal_address', ''),
        '{postal_address}': company_data.get('postal_address', ''),
        '{bank_account}': company_data.get('bank_account', ''),
        '{corr_account}': company_data.get('corr_account', ''),
        '{bank_name}': company_data.get('bank_name', ''),
        '{bik}': company_data.get('bik', ''),
        '{phone}': company_data.get('phone', ''),
        '{email}': company_data.get('email', ''),
        
        # Подписант
        '{signatory_full}': company_data.get('signatory_full', ''),
        '{signatory_short}': company_data.get('signatory_short', ''),
        '{signatory_position}': company_data.get('signatory_position', ''),
        '{signatory_position_genitive}': company_data.get('signatory_position_genitive', ''),
        '{signatory_position_dative}': company_data.get('signatory_position_dative', ''),
        '{signatory_position_accusative}': company_data.get('signatory_position_accusative', ''),
        '{signatory_position_instrumental}': company_data.get('signatory_position_instrumental', ''),
        '{signatory_position_prepositional}': company_data.get('signatory_position_prepositional', ''),
        '{signatory_gender}': 'женский' if company_data.get('signatory_gender') == 'feminine' else 'мужской',
        
        # Гендерно-согласованные формы
        '{based_on}': company_data.get('based_on', ''),
        '{based_on_agreed}': company_data.get('based_on_agreed', ''),
        '{based_on_agreed_full}': company_data.get('based_on_agreed_full', ''),
        '{acting_form}': company_data.get('acting_form', ''),
        '{acting_form_full}': company_data.get('acting_form_full', ''),
        
        # Другое
        '{edo_provider}': company_data.get('edo_provider', ''),
        '{domains}': company_data.get('domains', ''),
        
        # Договор
        '{contract_number}': company_data.get('contract_number', ''),
        '{contract_date}': company_data.get('contract_date_formatted', ''),
        '{contract_term_years}': company_data.get('contract_term_years', ''),
        '{contract_end_date}': company_data.get('contract_end_date_formatted', ''),
    }
    
    # Выполняем замену во всем документе
    replace_text_in_doc(doc, replacements)
    
    # Сохраняем результат
    try:
        doc.save(output_path)
    except Exception as e:
        raise Exception(f"Не удалось сохранить результат: {str(e)}")
    
    return output_path


def generate_contract_number():
    """
    Генерация номера договора
    """
    now = datetime.now()
    year = now.strftime('%Y')
    
    import random
    num = random.randint(1000, 9999)
    
    return f"СЛД-{num}{year}-М-{random.randint(10, 99)}"